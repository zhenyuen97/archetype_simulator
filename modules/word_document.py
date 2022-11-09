# modules
from docx import Document
from docx.opc.coreprops import CoreProperties
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

def save_word(text_list, name_of_company):
    logo = 'SUSS logo.png'
    company_archetypes = 'images/company_archetypes.png'
    emotion_chart = 'images/emotion_chart.png'
    employment_type = 'images/employment_type.png'
    former_employee_chart = 'images/former_employee_description.png'
    job_types_avg_score = 'images/job_types_avg_score.png'
    personality_chart = 'images/personality_chart.png'
    personality_type_avg_score = 'images/personality_type_avg_score.png'
    
    output = './report/Company Report.docx'

    document_description = f'ANL488 Report Summary ({name_of_company})' #eid.get() 

    comp_detail = '''Singapore University of Social Sciences
    463 Clementi Rd, Singapore 599494
    Email: zyng005@suss.edu.sg, Phone: +65 9837 3952
    '''

    # Certificate template
    body_text = f'''
The below summarises {name_of_company} employee information. The information obtained comes from Glassdoor.

Do note that the report is designed to help the company summarise employee description, and is for general information purposes only. All analyses on the site is provided in good faith and it make no representation of any kind, express or implied, regarding the accuracy, or completeness of the company.
                '''

    # create instance
    doc =  Document()

    # declare margin
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    section = doc.sections[0]

    # logo image placement
    logo = doc.add_picture(logo, width=Inches(2.52))
    logo_placement = doc.paragraphs[-1] 
    logo_placement.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # line space
    for _ in range(1):
        linespace_style = doc.styles['Body Text']
        linespace = doc.add_paragraph(style=linespace_style).add_run()
        linespace_style.font.size = Pt(10)
        

    # Header 
    heading_style = doc.styles['Body Text']
    head=doc.add_paragraph(style=heading_style).add_run(f'{document_description}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.font.size = Pt(14)
    head.font.bold = True 

    # body text 
    body_style = doc.styles['Body Text']
    body = doc.add_paragraph(style=body_style).add_run(f'{body_text}')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    body.font.size = Pt(14)
    body.font.name = 'Times New Roman'

    #line space
    for _ in range(2):
        linespace_style = doc.styles['Body Text']
        linespace = doc.add_paragraph(style=linespace_style).add_run()
        linespace.font.size = 10
    

    # signature image & text
    emotion_text = f'''The emotion chart illustrates the proportions of specific emotions (Anticipation, Trust, Joy, Surprise, Sadness, Fear, Anger and Disgust).

The calculation of emotion is based upon the NRC Emotion Lexicon, that contains a list of words and their associations with eight basic emotions. Each word in the employee review will be mapped against the NRC emotion lexicon to compute the overall emotion expressed by the employee.

Based on the chart above, {text_list[0]} are amongst the top 3 emotions expressed.
    '''
    emotion_description = doc.styles['Normal']
    doc.add_picture(emotion_chart, width = Inches(6))
    doc.add_paragraph(style=emotion_description).add_run(f'{emotion_text}')
    emotion_description.font.size = Pt(12)
    emotion_description.font.name = 'Arial'
    
    for _ in range(4):
        linespace_style = doc.styles['Body Text']
        linespace = doc.add_paragraph(style=linespace_style)
    
    # ================================
    # Personality Chart 
    # ================================
    personality_text = f'''The personality chart illustrates the proportions of 16 Myers-Briggs Personality Type (MBTI) Types. MBTI is a self-report inventory designed to identify a person's type, strengths, and preferences. Through a person's MBTI, we are able to determine his strengths, weaknesses, possible career preferences as well as their compatiblity with other people.

The calculation of personality is based upon the personality cafe, which is a forum that is dedicated to all ranges of personality types. The following derives a person's personality by applying text classification task using natural language processing and deep learning.

Based on the chart above, {text_list[1]} are amongst the top 3 personality type of {name_of_company} Employees.

{text_list[2]}

{text_list[3]}

{text_list[4]}
    '''
    personality_description_description = doc.styles['Normal']
    doc.add_picture(personality_chart, width = Inches(6))
    doc.add_paragraph(style=personality_description_description).add_run(f'{personality_text}')
    personality_description_description.font.size = Pt(12)
    personality_description_description.font.name = 'Arial'

    doc.add_page_break()

    # ================================
    # Employment Chart
    # ================================
    employment_text = f'''The employment chart illustrates the employment type of reviewers for {name_of_company} on Glassdoor.
    
The breakdown shows that the majority of reviewers are {text_list[5]}.
    '''
    employment_description = doc.styles['Normal']
    doc.add_picture(employment_type, width = Inches(6))
    doc.add_paragraph(style=employment_description).add_run(f'{employment_text}')
    employment_description.font.size = Pt(12)
    employment_description.font.name = 'Arial'
    
    doc.add_page_break()

    # ================================
    # Former Employee Chart
    # ================================
    former_employee_text = f'''A closer look at the duration of former employee gives us an overview of how often does an employee usually stays in the company.
    
It can be seen that the majority of employee stays for {text_list[6]}.
    '''
    former_employee_description = doc.styles['Normal']
    doc.add_picture(former_employee_chart, width = Inches(6))
    doc.add_paragraph(style=former_employee_description).add_run(f'{former_employee_text}')
    former_employee_description.font.size = Pt(12)
    former_employee_description.font.name = 'Arial'
    
    doc.add_page_break()

    # ================================
    # Company Archetype
    # ================================
    archetype_text = f'''The company archetype shows 3 archetypes of a company Innovative-Traditional, Open-Exclusive, and Extrovert-Introvert. 
    
Innovative and Traditional is derived by scraping the posts of innovative and traditional companies on LinkedIn, and using natural language processing (NLP) and Naive Bayes Classifier to classify reviewer's comments about the company.
    
Open and Exclusive is derived from the emotion, where a study suggests that a person whose review suggest anger, fear and disgust evaluates the company as exclusive and a person that shows joy, trust and anticipation on the company as exclusive.
    
Introvert and Extrovert defines the company personality, where the model is trained using posts by people of different personality from the personality cafe forum.
    
It can be seen that reviewers deemed {name_of_company} as {text_list[7]}, {text_list[8]} and {text_list[9]}
    '''
    archetype_description = doc.styles['Normal']
    doc.add_picture(company_archetypes, width = Inches(6))
    doc.add_paragraph(style=archetype_description).add_run(f'{archetype_text}')
    archetype_description.font.size = Pt(12)
    archetype_description.font.name = 'Arial'
    
    doc.add_page_break()
    #
    # ================================
    # Job type - Score
    # ================================
    job_type_text = f'''The job type chart gives us an overview of the top 5 roles that provide reviews for {name_of_company}. 
    
It also tells us the average score a person give on Glassdoor based on these positions. It can be seen that {text_list[10]}
    '''
    job_type_description = doc.styles['Normal']
    doc.add_picture(job_types_avg_score, width = Inches(6))
    doc.add_paragraph(style=job_type_description).add_run(f'{job_type_text}')
    job_type_description.font.size = Pt(12)
    job_type_description.font.name = 'Arial'
    
    doc.add_page_break()

    # ================================
    # Personality Type Score
    # ================================
    personality_type_text = f'''The personality tytpe breakdown shows us the average score given by the different personality on the company. This allow us to have a better sense as to which personality type would be more satisfied working at {name_of_company} as compared to the others. 
    
{text_list[11]}

{text_list[12]}
    '''
    personality_type_description = doc.styles['Normal']
    doc.add_picture(personality_type_avg_score, width = Inches(6))
    doc.add_paragraph(style=personality_type_description).add_run(f'{personality_type_text}')
    personality_type_description.font.size = Pt(12)
    personality_type_description.font.name = 'Arial'
    
    doc.add_page_break()

    # footer text : company description
    company_text = doc.styles['Normal']
    company_text.paragraph_format.space_before = Pt(12)
    doc.add_paragraph(style=company_text).add_run(f'{comp_detail}')
    center_align = doc.paragraphs[-1] 
    center_align.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # saving file to word document
    doc.save(output)
